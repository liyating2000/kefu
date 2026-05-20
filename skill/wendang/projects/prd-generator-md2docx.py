#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Markdown转Word文档转换脚本（增强版）
"""

import argparse
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    设置单元格边框
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def main():
    parser = argparse.ArgumentParser(description='Markdown转Word文档')
    parser.add_argument('--input-file', required=True, help='Markdown文件路径')
    parser.add_argument('--output-file', required=True, help='输出Word文件路径')
    args = parser.parse_args()

    try:
        # 读取Markdown文件
        with open(args.input_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        # 创建Word文档
        doc = Document()

        # 设置默认字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        doc.styles['Normal'].font.size = Pt(12)

        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

        # 解析Markdown内容
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()

            # 跳过空行
            if not line:
                i += 1
                continue

            # 一级标题
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p.runs[0].font.name = '黑体'
                p.runs[0].font.size = Pt(18)
                p.runs[0].font.color.rgb = RGBColor(0, 0, 0)

            # 二级标题
            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)
                p.runs[0].font.name = '黑体'
                p.runs[0].font.size = Pt(16)
                p.runs[0].font.bold = True

            # 三级标题
            elif line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)
                p.runs[0].font.name = '黑体'
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.bold = True

            # 四级标题
            elif line.startswith('#### '):
                p = doc.add_heading(line[5:], level=4)
                p.runs[0].font.name = '黑体'
                p.runs[0].font.size = Pt(13)
                p.runs[0].font.bold = True

            # 列表项
            elif line.startswith('- '):
                p = doc.add_paragraph(line[2:], style='List Bullet')

            # 数字列表
            elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
                text = line[3:]
                p = doc.add_paragraph(text, style='List Number')

            # 表格行（Markdown表格）
            elif line.startswith('|') and line.endswith('|'):
                # 收集表格所有行
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].rstrip())
                    i += 1
                i -= 1  # 回退一行，因为外层循环会+1

                # 解析表格
                table_data = []
                for table_line in table_lines:
                    cells = [cell.strip() for cell in table_line.split('|')[1:-1]]
                    # 跳过分隔线
                    if all(cell.replace('-', '').replace(':', '').strip() == '' for cell in cells):
                        continue
                    table_data.append(cells)

                if table_data:
                    # 创建表格
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Table Grid'
                    table.autofit = False

                    # 设置表格列宽
                    for row in table.rows:
                        for cell in row.cells:
                            cell.width = Inches(1.5)

                    # 填充表格数据
                    for row_idx, row_data in enumerate(table_data):
                        row = table.rows[row_idx]
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < len(row.cells):
                                cell = row.cells[col_idx]
                                cell.text = cell_data

                                # 设置表头样式
                                if row_idx == 0:
                                    cell.paragraphs[0].runs[0].font.name = '黑体'
                                    cell.paragraphs[0].runs[0].font.size = Pt(12)
                                    cell.paragraphs[0].runs[0].font.bold = True
                                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    # 设置表头背景色
                                    shading_elm = OxmlElement('w:shd')
                                    shading_elm.set(qn('w:fill'), 'D9D9D9')
                                    cell._tc.get_or_add_tcPr().append(shading_elm)
                                else:
                                    cell.paragraphs[0].runs[0].font.name = '宋体'
                                    cell.paragraphs[0].runs[0].font.size = Pt(11)
                                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # 普通段落
            else:
                p = doc.add_paragraph(line)
                p.paragraph_format.line_spacing = 1.5

            i += 1

        # 保存Word文档
        doc.save(args.output_file)

        result = {
            "status": "success",
            "message": "Word文档生成成功",
            "input_file": args.input_file,
            "output_file": args.output_file
        }

    except Exception as e:
        import traceback
        result = {
            "status": "error",
            "message": str(e),
            "traceback": traceback.format_exc(),
            "input_file": args.input_file,
            "output_file": args.output_file
        }

    print(json.dumps(result, ensure_ascii=False))

if __name__ == "__main__":
    main()
